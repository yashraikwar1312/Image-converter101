from flask import Flask, request, jsonify, send_file, render_template_string
from werkzeug.utils import secure_filename
import os
import io
import zipfile
from datetime import datetime
import mimetypes

# Image processing
from PIL import Image, ImageDraw
import cairosvg

# Document processing
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

# Data processing
import pandas as pd
import json
import xml.etree.ElementTree as ET
import yaml
import csv

# Additional libraries
import base64
import tempfile
import shutil
import logging

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Create directories
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

ALLOWED_EXTENSIONS = {
    'png', 'jpg', 'jpeg', 'gif', 'bmp', 'svg', 'webp',
    'pdf', 'docx', 'doc', 'txt', 'rtf',
    'xlsx', 'xls', 'csv', 'ods',
    'json', 'xml', 'yaml', 'yml'
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_extension(filename):
    return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

def generate_output_filename(original_filename, target_format):
    name = os.path.splitext(original_filename)[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{name}_{timestamp}.{target_format}"

class FileConverter:
    @staticmethod
    def convert_image(input_path, output_path, target_format):
        """Convert image files between different formats"""
        try:
            if target_format.lower() == 'svg':
                # Special handling for SVG conversion
                return FileConverter.convert_to_svg(input_path, output_path)
            
            with Image.open(input_path) as img:
                # Handle transparency for formats that don't support it
                if target_format.lower() in ['jpg', 'jpeg'] and img.mode in ['RGBA', 'LA']:
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                
                # Convert and save
                if target_format.lower() == 'webp':
                    img.save(output_path, 'WEBP', quality=90)
                else:
                    img.save(output_path, target_format.upper())
            
            return True
        except Exception as e:
            logger.error(f"Image conversion error: {e}")
            return False

    @staticmethod
    def convert_to_svg(input_path, output_path):
        """Convert image to SVG format"""
        try:
            with Image.open(input_path) as img:
                width, height = img.size
                
                # Convert image to base64
                buffered = io.BytesIO()
                img.save(buffered, format="PNG")
                img_str = base64.b64encode(buffered.getvalue()).decode()
                
                # Create SVG with embedded image
                svg_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg width="{width}" height="{height}" xmlns="http://www.w3.org/2000/svg">
    <image href="data:image/png;base64,{img_str}" width="{width}" height="{height}"/>
</svg>'''
                
                with open(output_path, 'w') as f:
                    f.write(svg_content)
            
            return True
        except Exception as e:
            logger.error(f"SVG conversion error: {e}")
            return False

    @staticmethod
    def convert_to_pdf(input_path, output_path, source_format):
        """Convert various formats to PDF"""
        try:
            if source_format in ['jpg', 'jpeg', 'png', 'bmp', 'gif']:
                return FileConverter.image_to_pdf(input_path, output_path)
            elif source_format in ['txt']:
                return FileConverter.text_to_pdf(input_path, output_path)
            elif source_format in ['csv']:
                return FileConverter.csv_to_pdf(input_path, output_path)
            else:
                return False
        except Exception as e:
            logger.error(f"PDF conversion error: {e}")
            return False

    @staticmethod
    def image_to_pdf(input_path, output_path):
        """Convert image to PDF"""
        try:
            with Image.open(input_path) as img:
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                img.save(output_path, 'PDF', resolution=100.0)
            return True
        except Exception as e:
            logger.error(f"Image to PDF error: {e}")
            return False

    @staticmethod
    def text_to_pdf(input_path, output_path):
        """Convert text file to PDF"""
        try:
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            with open(input_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            y = height - 50
            for line in lines:
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, line.strip())
                y -= 20
            
            c.save()
            return True
        except Exception as e:
            logger.error(f"Text to PDF error: {e}")
            return False

    @staticmethod
    def csv_to_pdf(input_path, output_path):
        """Convert CSV to PDF"""
        try:
            df = pd.read_csv(input_path)
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # Write headers
            y = height - 50
            x = 50
            for col in df.columns:
                c.drawString(x, y, str(col))
                x += 100
            
            y -= 30
            
            # Write data
            for _, row in df.iterrows():
                if y < 50:
                    c.showPage()
                    y = height - 50
                
                x = 50
                for value in row:
                    c.drawString(x, y, str(value)[:15])  # Truncate long values
                    x += 100
                y -= 20
            
            c.save()
            return True
        except Exception as e:
            logger.error(f"CSV to PDF error: {e}")
            return False

    @staticmethod
    def pdf_to_image(input_path, output_path, target_format):
        """Convert PDF to image"""
        try:
            doc = fitz.open(input_path)
            page = doc.load_page(0)  # Convert first page
            
            # Render page to image
            mat = fitz.Matrix(2, 2)  # 2x zoom
            pix = page.get_pixmap(matrix=mat)
            
            if target_format.lower() == 'png':
                pix.save(output_path)
            else:
                # Convert to PIL Image for other formats
                img_data = pix.tobytes("ppm")
                img = Image.open(io.BytesIO(img_data))
                img.save(output_path, target_format.upper())
            
            doc.close()
            return True
        except Exception as e:
            logger.error(f"PDF to image error: {e}")
            return False

    @staticmethod
    def convert_document(input_path, output_path, source_format, target_format):
        """Convert between document formats"""
        try:
            if source_format == 'txt' and target_format == 'docx':
                return FileConverter.txt_to_docx(input_path, output_path)
            elif source_format == 'docx' and target_format == 'txt':
                return FileConverter.docx_to_txt(input_path, output_path)
            else:
                return False
        except Exception as e:
            logger.error(f"Document conversion error: {e}")
            return False

    @staticmethod
    def txt_to_docx(input_path, output_path):
        """Convert text file to DOCX"""
        try:
            doc = Document()
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc.add_paragraph(content)
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"TXT to DOCX error: {e}")
            return False

    @staticmethod
    def docx_to_txt(input_path, output_path):
        """Convert DOCX to text file"""
        try:
            doc = Document(input_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text))
            return True
        except Exception as e:
            logger.error(f"DOCX to TXT error: {e}")
            return False

    @staticmethod
    def convert_data_format(input_path, output_path, source_format, target_format):
        """Convert between data formats (CSV, JSON, XML, YAML)"""
        try:
            if source_format == 'csv':
                df = pd.read_csv(input_path)
                data = df.to_dict('records')
            elif source_format == 'json':
                with open(input_path, 'r') as f:
                    data = json.load(f)
            elif source_format in ['yaml', 'yml']:
                with open(input_path, 'r') as f:
                    data = yaml.safe_load(f)
            elif source_format == 'xml':
                tree = ET.parse(input_path)
                root = tree.getroot()
                data = FileConverter.xml_to_dict(root)
            else:
                return False

            # Convert to target format
            if target_format == 'csv':
                if isinstance(data, list):
                    df = pd.DataFrame(data)
                    df.to_csv(output_path, index=False)
                else:
                    return False
            elif target_format == 'json':
                with open(output_path, 'w') as f:
                    json.dump(data, f, indent=2)
            elif target_format == 'yaml':
                with open(output_path, 'w') as f:
                    yaml.dump(data, f, default_flow_style=False)
            elif target_format == 'xml':
                root = ET.Element('root')
                FileConverter.dict_to_xml(data, root)
                tree = ET.ElementTree(root)
                tree.write(output_path, encoding='utf-8', xml_declaration=True)
            else:
                return False

            return True
        except Exception as e:
            logger.error(f"Data format conversion error: {e}")
            return False

    @staticmethod
    def xml_to_dict(element):
        """Convert XML element to dictionary"""
        result = {}
        for child in element:
            if len(child) == 0:
                result[child.tag] = child.text
            else:
                result[child.tag] = FileConverter.xml_to_dict(child)
        return result

    @staticmethod
    def dict_to_xml(data, parent):
        """Convert dictionary to XML elements"""
        if isinstance(data, dict):
            for key, value in data.items():
                elem = ET.SubElement(parent, str(key))
                if isinstance(value, (dict, list)):
                    FileConverter.dict_to_xml(value, elem)
                else:
                    elem.text = str(value)
        elif isinstance(data, list):
            for i, item in enumerate(data):
                elem = ET.SubElement(parent, f'item_{i}')
                FileConverter.dict_to_xml(item, elem)

    @staticmethod
    def convert_spreadsheet(input_path, output_path, source_format, target_format):
        """Convert between spreadsheet formats"""
        try:
            # Read the spreadsheet
            if source_format == 'csv':
                df = pd.read_csv(input_path)
            elif source_format in ['xlsx', 'xls']:
                df = pd.read_excel(input_path)
            elif source_format == 'ods':
                df = pd.read_excel(input_path, engine='odf')
            else:
                return False

            # Write to target format
            if target_format == 'csv':
                df.to_csv(output_path, index=False)
            elif target_format == 'xlsx':
                df.to_excel(output_path, index=False, engine='openpyxl')
            elif target_format == 'ods':
                df.to_excel(output_path, index=False, engine='odf')
            else:
                return False

            return True
        except Exception as e:
            logger.error(f"Spreadsheet conversion error: {e}")
            return False

def perform_conversion(input_path, output_path, source_format, target_format):
    """Main conversion function that routes to appropriate converter"""
    try:
        source_format = source_format.lower()
        target_format = target_format.lower()

        # Image conversions
        if source_format in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'svg', 'webp']:
            if target_format in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'svg', 'webp']:
                return FileConverter.convert_image(input_path, output_path, target_format)
            elif target_format == 'pdf':
                return FileConverter.convert_to_pdf(input_path, output_path, source_format)

        # PDF conversions
        elif source_format == 'pdf':
            if target_format in ['png', 'jpg', 'jpeg', 'bmp', 'gif']:
                return FileConverter.pdf_to_image(input_path, output_path, target_format)

        # Document conversions
        elif source_format in ['txt', 'docx', 'doc']:
            if target_format == 'pdf':
                return FileConverter.convert_to_pdf(input_path, output_path, source_format)
            elif target_format in ['txt', 'docx']:
                return FileConverter.convert_document(input_path, output_path, source_format, target_format)

        # Data format conversions
        elif source_format in ['csv', 'json', 'xml', 'yaml', 'yml']:
            if target_format in ['csv', 'json', 'xml', 'yaml']:
                return FileConverter.convert_data_format(input_path, output_path, source_format, target_format)
            elif target_format == 'pdf':
                return FileConverter.convert_to_pdf(input_path, output_path, source_format)

        # Spreadsheet conversions
        elif source_format in ['xlsx', 'xls', 'ods']:
            if target_format in ['csv', 'xlsx', 'ods']:
                return FileConverter.convert_spreadsheet(input_path, output_path, source_format, target_format)

        return False
    except Exception as e:
        logger.error(f"Conversion error: {e}")
        return False

@app.route('/')
def index():
    """Serve the main HTML page"""
    try:
        with open('index.html', 'r') as f:
            return f.read()
    except FileNotFoundError:
        # Return a basic HTML page if index.html doesn't exist
        return '''
        <!DOCTYPE html>
        <html>
        <head>
            <title>Universal File Converter</title>
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body>
            <h1>Universal File Converter</h1>
            <p>Please create an index.html file with the frontend code.</p>
            <p>API endpoint: POST /convert</p>
        </body>
        </html>
        '''

@app.route('/convert', methods=['POST'])
def convert_file():
    """Handle file conversion requests"""
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        target_format = request.form.get('format', '').lower()

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not target_format:
            return jsonify({'error': 'No target format specified'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not supported'}), 400

        # Secure the filename
        original_filename = secure_filename(file.filename)
        source_format = get_file_extension(original_filename)

        # Generate unique filenames
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        input_filename = f"{timestamp}_{original_filename}"
        output_filename = generate_output_filename(original_filename, target_format)

        # Save uploaded file
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], input_filename)
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

        file.save(input_path)

        # Perform conversion
        success = perform_conversion(input_path, output_path, source_format, target_format)

        if success and os.path.exists(output_path):
            # Clean up input file
            try:
                os.remove(input_path)
            except:
                pass

            # Send the converted file
            return send_file(
                output_path,
                as_attachment=True,
                download_name=f"converted.{target_format}",
                mimetype=mimetypes.guess_type(output_path)[0]
            )
        else:
            # Clean up files
            try:
                os.remove(input_path)
                if os.path.exists(output_path):
                    os.remove(output_path)
            except:
                pass

            return jsonify({'error': 'Conversion failed'}), 500

    except Exception as e:
        logger.error(f"Conversion endpoint error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/health')
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'timestamp': datetime.now().isoformat()})

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large'}), 413

@app.errorhandler(500)
def internal_error(e):
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == '__main__':
    # Create necessary directories
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    
    # Run the application
    app.run(debug=True, host='0.0.0.0', port=5000)
            